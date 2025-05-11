# job_ad_generator_project/content/predefined_data.py
import os
try:
    import docx
    from docx.document import Document as _Document # To access block items
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    docx = None
    _Document = None
    CT_P = None
    CT_Tbl = None
    _Cell = None
    Table = None
    Paragraph = None
    print("WARNING: 'python-docx' library not found. .docx file support will be disabled.")
    print("Please install it by running: pip install python-docx")

from configs.app_settings import AD_TEMPLATES_DIR, JD_DESCRIPTIONS_DIR # Assuming this is correc

# --- Default Base Strings (Built-in) ---
DEFAULT_JOB_AD_TEMPLATE = """
**Job Title:** [Insert Job Title Here]
**Company:** [Your Company Name]
**Location:** [Location - e.g., Remote, City, State]

**About Us:**
[Provide a brief, engaging description of your company culture, mission, and values. What makes your company a great place to work?]

**Job Summary:**
[Briefly (2-3 sentences) describe the main purpose and essence of this role. What will this person achieve?]

**Key Responsibilities:**
*   [Responsibility 1: Action-oriented, e.g., "Develop and maintain..."]
*   [Responsibility 2: Be specific]
*   [Responsibility 3]
*   ...

**Qualifications & Skills:**
*   **Required:**
    *   [Qualification 1: e.g., Bachelor's degree in Computer Science or equivalent experience]
    *   [Skill 1: e.g., X+ years of experience in Y technology]
    *   [Skill 2: e.g., Proficiency in Z software/language]
    *   ...
*   **Preferred (Bonus Points):**
    *   [Optional Skill 1: e.g., Experience with A]
    *   [Optional Certification: e.g., Certification B]
    *   ...

**Why Join Us? (Benefits & Perks):**
*   [Benefit 1: e.g., Competitive salary and comprehensive benefits package (health, dental, vision)]
*   [Benefit 2: e.g., Opportunities for professional growth and development]
*   [Benefit 3: e.g., Flexible work arrangements / Remote options]
*   [Benefit 4: e.g., A collaborative, innovative, and inclusive work environment]
*   ...

**How to Apply:**
[Clear instructions on how to apply. e.g., "Interested candidates are invited to submit their resume and a cover letter outlining their suitability for the role to [email_address] / via our careers page: [Link]" Be specific about what you need.]

**Equal Opportunity Employer Statement:**
[Your Company Name] is an equal opportunity employer. We celebrate diversity and are committed to creating an inclusive environment for all employees.
""".strip()

DEFAULT_JOB_DESCRIPTION = """
Role: Senior Software Engineer (Backend Focus)
Team: Core Platform Engineering
Reports to: Engineering Manager

We are seeking a highly skilled Senior Software Engineer to join our innovative Core Platform team.
This individual will play a key role in designing, developing, and deploying robust and scalable backend systems that power our flagship products.
The ideal candidate is passionate about building high-performance services, enjoys tackling complex technical challenges, and thrives in a collaborative, fast-paced environment.

Primary Responsibilities:
- Architect and implement new microservices and APIs.
- Optimize existing backend systems for performance, scalability, and reliability.
- Collaborate with frontend developers, product managers, and other stakeholders to deliver new features.
- Write clean, maintainable, and well-tested code (Python, Go).
- Mentor junior engineers and contribute to best practices within the team.
- Participate in code reviews and design discussions.
- Troubleshoot and resolve production issues.

Required Qualifications:
- 5+ years of professional software development experience with a focus on backend systems.
- Strong proficiency in Python and/or Go.
- Extensive experience with designing and building RESTful APIs and microservices.
- Solid understanding of database technologies (e.g., PostgreSQL, MongoDB, Cassandra).
- Experience with cloud platforms (AWS, GCP, or Azure), including serverless and containerization (Docker, Kubernetes).
- Familiarity with message queues (e.g., Kafka, RabbitMQ) and caching mechanisms (e.g., Redis).
- Excellent problem-solving and analytical skills.
- Bachelor's degree in Computer Science or a related field, or equivalent practical experience.

Preferred Qualifications:
- Experience with gRPC.
- Knowledge of data streaming technologies.
- Contributions to open-source projects.
""".strip()


# --- Helper Function to Read .docx Files (Improved for Tables) ---
def _read_docx_file(filepath):
    if docx is None or _Document is None: # Check all necessary imports
        print(f"Skipping .docx file {filepath} as python-docx components are not available.")
        return ""
    
    try:
        document = docx.Document(filepath)
        full_text_parts = []

        # Iterate through document body elements (paragraphs and tables in order)
        # This requires accessing the underlying XML structure a bit for docx.Document.element.body
        # A more direct way if available in your python-docx version is preferred,
        # but iterating through doc.element.body is a common way to get mixed content.

        # Simpler approach for iterating through block items if directly supported:
        # (This assumes a version of python-docx where document.iter_block_items() or similar exists)
        # However, iterating doc.element.body is more universally compatible for older versions too.
        
        # Let's try iterating through a known sequence of paragraphs and tables.
        # If you need strict order of mixed content, a more complex iteration over doc.element.body is needed.
        # For now, extract all paragraphs, then all tables.
        # This might not preserve perfect interleaving if a table is between two paragraphs.

        # Extract all paragraph text first
        for para in document.paragraphs:
            full_text_parts.append(para.text)
        
        full_text_parts.append("\n\n--- TABLES DATA ---\n") # Separator

        # Extract all table text
        for table in document.tables:
            full_text_parts.append(f"\nTable (Rows: {len(table.rows)}, Columns: {len(table.columns)}):\n")
            for i, row in enumerate(table.rows):
                row_texts = []
                for j, cell in enumerate(row.cells):
                    # Clean cell text: replace newlines within a cell with spaces for better LLM processing
                    cell_text = cell.text.replace('\n', ' ').strip()
                    row_texts.append(cell_text)
                full_text_parts.append(f"Row {i+1}: | " + " | ".join(row_texts) + " |")
            full_text_parts.append("\n--- END OF TABLE ---")

        return '\n'.join(full_text_parts).strip()

    except Exception as e:
        print(f"Error reading .docx file {filepath}: {e}")
        return ""

# --- Iterating through body elements for strict order (More Advanced) ---
# This is a more robust way to get content in document order if you need precise interleaving.
def _read_docx_file_ordered(filepath):
    if docx is None or _Document is None or CT_P is None or CT_Tbl is None:
        print(f"Skipping .docx file {filepath} as python-docx components for ordered reading are not available.")
        return ""
    try:
        document = docx.Document(filepath)
        full_text_parts = []

        for block in document.element.body:
            if isinstance(block, CT_P): # Paragraph
                # Reconstruct Paragraph object to use its .text property
                p = Paragraph(block, document)
                full_text_parts.append(p.text)
            elif isinstance(block, CT_Tbl): # Table
                # Reconstruct Table object
                table = Table(block, document)
                full_text_parts.append(f"\n\n--- TABLE START ---\nTable (Rows: {len(table.rows)}, Columns: {len(table.columns)}):\n")
                for i, row in enumerate(table.rows):
                    row_texts = []
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.replace('\n', ' ').strip()
                        row_texts.append(f"Cell({i+1},{j+1}): {cell_text}") # More descriptive
                    full_text_parts.append(" | ".join(row_texts)) # Join cells for the row
                full_text_parts.append("--- TABLE END ---\n")
        
        return '\n'.join(full_text_parts).strip()
    except Exception as e:
        print(f"Error reading .docx file (ordered) {filepath}: {e}")
        return ""


# --- Helper Function to Load Content from Files ---
# Choose which docx reader to use: _read_docx_file (simpler) or _read_docx_file_ordered (more complex but better order)
DOCX_READER_FUNCTION = _read_docx_file_ordered # Or _read_docx_file

def _load_content_from_directory(directory_path, allowed_extensions=(".txt", ".md", ".docx")):
    """
    Scans a directory for files with allowed extensions and reads their content.
    The filename (without extension) becomes the key, and file content the value.
    """
    loaded_content = {}
    if not os.path.isdir(directory_path):
        return loaded_content

    for filename in os.listdir(directory_path):
        name_lower = filename.lower()
        filepath = os.path.join(directory_path, filename)
        content = ""

        if name_lower.endswith((".txt", ".md")): # Combined check for .txt and .md
            if ".txt" in allowed_extensions or ".md" in allowed_extensions: # Ensure it's allowed
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                except Exception as e:
                    print(f"Error loading text/md file {filepath}: {e}")
                    continue
        elif name_lower.endswith(".docx"):
            if ".docx" in allowed_extensions:
                content = DOCX_READER_FUNCTION(filepath) # Use the chosen reader function
        else:
            continue

        if content:
            key_name = os.path.splitext(filename)[0].replace("_", " ").replace("-", " ").title()
            loaded_content[key_name] = content
            
    return loaded_content

# --- Initialize Dictionaries for Presets ---
PREDEFINED_TEMPLATES = {
    "Default Modern Template": DEFAULT_JOB_AD_TEMPLATE,
}

PREDEFINED_DESCRIPTIONS = {
    "Senior Software Engineer (Backend)": DEFAULT_JOB_DESCRIPTION,
}

# --- Load and Merge from Files ---
# Load custom ad templates from files
# --- Load and Merge from Files ---
custom_ad_templates = _load_content_from_directory(AD_TEMPLATES_DIR)
PREDEFINED_TEMPLATES.update(custom_ad_templates)

custom_jd_descriptions = _load_content_from_directory(JD_DESCRIPTIONS_DIR)
PREDEFINED_DESCRIPTIONS.update(custom_jd_descriptions)

# Optional Debugging
# print("--- Loaded PREDEFINED_TEMPLATES (including files) ---")
# for k in PREDEFINED_TEMPLATES.keys(): print(f"Key: {k}")
# print("--- Loaded PREDEFINED_DESCRIPTIONS (including files) ---")
# for k in PREDEFINED_DESCRIPTIONS.keys(): print(f"Key: {k}")